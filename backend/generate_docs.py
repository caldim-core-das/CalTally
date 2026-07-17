import os
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("python-docx not installed. Run 'pip install python-docx'")
    exit(1)

doc = Document()

# Title
title = doc.add_heading('Complete Testing Guide - CalBooks (Tally Clone)', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('This document provides a comprehensive end-to-end testing workflow to ensure that the core accounting modules (Ledgers, Vouchers, P&L, Balance Sheet, Trial Balance, and Cash Flow) are functioning correctly mathematically and logically.')

# Section 1
doc.add_heading('1. Initial Setup (Ledgers & Opening Balances)', level=1)
doc.add_paragraph('Test the creation of basic ledgers and the golden rule of double-entry accounting for opening balances.')
p = doc.add_paragraph()
p.add_run('Steps:\n').bold = True
p.add_run('1. Go to Accountant Tools -> Ledgers.\n')
p.add_run('2. Verify the "Capital Account" exists (or create it under Capital Account group). Give it an Opening Balance of ₹7,00,000 (Cr).\n')
p.add_run('3. Create a Bank Account named "SBI Bank Account" under Bank Accounts group. Give it an Opening Balance of ₹5,00,000 (Dr).\n')
p.add_run('4. Create a Bank Account named "HDFC Bank Account" under Bank Accounts group. Give it an Opening Balance of ₹2,00,000 (Dr).\n')
p = doc.add_paragraph()
p.add_run('Expected Result:\n').bold = True
p.add_run('The total Debit opening balances (5L + 2L) exactly match the total Credit opening balances (7L). The system should report no imbalances.')

# Section 2
doc.add_heading('2. Transaction Entry (Vouchers)', level=1)
doc.add_paragraph('Test the recording of day-to-day business expenses.')
p = doc.add_paragraph()
p.add_run('Steps:\n').bold = True
p.add_run('1. Create a Ledger named "Rent Expense" under Indirect Expenses.\n')
p.add_run('2. Go to Vouchers and create a Payment Voucher.\n')
p.add_run('3. Debit "Rent Expense" for ₹25,000.\n')
p.add_run('4. Credit "SBI Bank Account" for ₹25,000.\n')
p.add_run('5. Save the voucher.')

# Section 3
doc.add_heading('3. Verifying Profit & Loss (P&L)', level=1)
doc.add_paragraph('Test that expenses correctly affect the company\'s bottom line.')
p = doc.add_paragraph()
p.add_run('Steps:\n').bold = True
p.add_run('1. Navigate to Reports -> Profit & Loss.\n')
p.add_run('2. Check the Indirect Expenses section.\n')
p = doc.add_paragraph()
p.add_run('Expected Result:\n').bold = True
p.add_run('Rent Expense should show ₹25,000. The Net Profit/Loss at the bottom should display a Net Loss of -₹25,000 (since no income has been recorded yet).')

# Section 4
doc.add_heading('4. Verifying Balance Sheet', level=1)
doc.add_paragraph('Test the fundamental accounting equation: Assets = Liabilities + Equity.')
p = doc.add_paragraph()
p.add_run('Steps:\n').bold = True
p.add_run('1. Navigate to Reports -> Balance Sheet.\n')
p.add_run('2. Review the Equity & Liabilities (Left Side) and Assets (Right Side).\n')
p = doc.add_paragraph()
p.add_run('Expected Result:\n').bold = True
p.add_run('Left Side:\n- Capital Account: ₹7,00,000\n- Profit & Loss (Net Loss): -₹25,000\n- Total Liabilities: ₹6,75,000\n\n')
p.add_run('Right Side:\n- Current Assets (Bank Accounts): ₹6,75,000 (SBI is now ₹4,75,000, HDFC is ₹2,00,000)\n- Total Assets: ₹6,75,000\n\nThe Balance Sheet should balance perfectly with no red warnings.')

# Section 5
doc.add_heading('5. Verifying Trial Balance', level=1)
doc.add_paragraph('Test that all ledger balances sum up correctly.')
p = doc.add_paragraph()
p.add_run('Steps:\n').bold = True
p.add_run('1. Navigate to Reports -> Trial Balance.\n')
p.add_run('2. Review the Grand Totals at the bottom of the table.\n')
p = doc.add_paragraph()
p.add_run('Expected Result:\n').bold = True
p.add_run('Total Closing Debits (SBI + HDFC + Rent) should equal ₹7,00,000. Total Closing Credits (Capital) should equal ₹7,00,000. The top status card should say "Balanced".')

# Section 6
doc.add_heading('6. Verifying Cash Flow Statement', level=1)
doc.add_paragraph('Test that the indirect cash flow accurately tracks cash movements.')
p = doc.add_paragraph()
p.add_run('Steps:\n').bold = True
p.add_run('1. Navigate to Reports -> Cash Flow.\n')
p.add_run('2. Check the Net Cash Flow figure.\n')
p = doc.add_paragraph()
p.add_run('Expected Result:\n').bold = True
p.add_run('Net Cash Flow should be +₹6,75,000. This is comprised of ₹7,00,000 (Cash from Financing/Capital) and -₹25,000 (Cash from Operations/Rent).')

# Save to the root of the project
doc.save('../Testing_Guide.docx')
print("Document saved successfully as Testing_Guide.docx in the project root.")
